from docx import Document
from docx.shared import Cm
import os
import pdfkit
from fpdf import FPDF

document = Document()

document.add_heading('List of highest individual scores in ODIs',0)

records = [
    [1,'Rohit Sharma','Rohit_Gurunath_Sharma.jpg',"Sharma is an aggressive batsman but plays with style and elegance. He is usually an opening batsman in limited overs cricket, but has played most of his Test cricket as a middle-order batsman.", 264, 33, 9, 152.60],
    [2,'Martin Guptill','Martin_Guptill_2_(cropped).jpg',"Guptill is the first cricketer from New Zealand  and holds the current record for the highest individual score in Cricket World Cup matches and the second highest score in One Day Internationals of 237 not out.", 237, 24, 11, 145.39],
    [3, 'Virender Sehwag', 'images.jpeg',"Virender Sehwag is a former Indian cricketer and former captain of the Indian National Cricket Team. He was renowned as a destructive and aggressive player. He holds numerous records in international and Indian cricket. He is a right arm batsman and right arm off-spin bowler.", 219, 25, 7, 146.97],
    [4, 'Chris Gayle', '6b0aaf92-a634-4164-b41a-134ae35c20d1.original.jpg',"A damaging allrounder, Gayle also claimed 158 wickets. He has 22 centuries, 47 fifties and a five-wicket haul to his name. His aggressive style is well suited to limited-overs cricket,", 215, 10, 16, 146.25],
    [5, 'Fakhar Zaman', 'Fakhar_Zaman,_Pakistan_vs_Sri_Lanka,_1st_ODI,_2017.jpg',"A left-handed opener with an unusually high backlift, Zaman relies on timing for his runs, which he's made plenty of. This is illustrated by his domestic average (42 in first class cricket and 49 in List A). ", 210, 24, 5, 134.61],
    [6, 'Pathum Nissanka', 'gettyimages-1498885383-612x612.jpg',"He is second highest run scorer for Sri Lanka scoring 332 runs in 9 matches with average of 41.5. On 9 February 2024, during the first ODI match against Afghanistan, Nissanka became the first Sri Lankan to score an ODI double hundred", 210, 20, 8, 151.07],
    [7, 'Ishan Kishan', 'Ishan-Kishan_64b4ea6bb98bb.png',"Ishan Kishan (born 18 July 1998) is an Indian international cricketer who plays for the Indian national cricket team as a wicket-keeper batsman. He made his international debut in March 2021 against England.", 210, 24, 10, 160.30],
    [8, 'Rohit Sharma', 'Rohit_Gurunath_Sharma.jpg', "Sharma is an aggressive batsman but plays with style and elegance. He is usually an opening batsman in limited overs cricket, but has played most of his Test cricket as a middle-order batsman.",209, 12, 16, 132.28],
    [9, 'Rohit Sharma', 'Rohit_Gurunath_Sharma.jpg',"Sharma is an aggressive batsman but plays with style and elegance. He is usually an opening batsman in limited overs cricket, but has played most of his Test cricket as a middle-order batsman.", 208, 13, 12, 135.94],
    [10, 'Shubman Gill', 'shubman-gill.png',"Shubman Gill is a talented top-order batter who hails from Punjab. He grew up playing cricket in his farm, and later, his father shifted his base to Mohali and made Gill join the Punjab Cricket Association.", 208, 19, 9, 139.59],
]


ODIstable = document.add_table(rows=1, cols=5)

# ODIstable.style = 'Table Grid'

ODIstable.style = 'Table Grid'
hdr_Cells = ODIstable.rows[0].cells
hdr_Cells[0].text = 'Sl no'
hdr_Cells[1].text = 'Name'
hdr_Cells[2].text = 'Image'
hdr_Cells[3].text = 'Description'
hdr_Cells[4].text = 'Score'
# hdr_Cells[5].text = '4s'
# hdr_Cells[6].text = '6s'
# hdr_Cells[7].text = 'Strike rate'

for Rank,Name,Image,Description,Score,fours,sixes,Strikerate in records:
    row_Cells = ODIstable.add_row().cells
    row_Cells[0].text = str(Rank)
    row_Cells[1].text = Name
    image_path = os.path.join("batsman_images", Image)
    if os.path.exists(image_path):
        paragraph = row_Cells[2].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Cm(3), height=Cm(3.5))
    else:
        row_Cells[2].text = "Image not found"
    row_Cells[3].text = Description
    row_Cells[4].text = str(Score)
    # row_Cells[5].text = str(fours)
    # row_Cells[6].text = str(sixes)
    # row_Cells[7].text = str(Strikerate)
    for col_idx, col in enumerate(ODIstable.columns):
        if col_idx == 0:
            col.width = Cm(1)
        if col_idx == 2:
            col.width = Cm(4)
        if col_idx == 3:
            col.width = Cm(7)
        if col_idx == 4:
            col.width = Cm(1.7)

# # open in word
# document.save('Table_scores.docx')
# os.system("xdg-open Table_scores.docx")

# open in pdf
word_file = 'Table_scores.docx'
document.save(word_file)
pdf_file = 'Table_scores.pdf'
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)

docx_to_pdf_command = f'libreoffice --headless --convert-to pdf {word_file}'
os.system(docx_to_pdf_command)

os.rename('Table_scores.pdf', pdf_file)
os.remove(word_file)
os.system(f"xdg-open {pdf_file}")



